"""
营销专业基本情况 Word 生成器
替换模版中"一、电力电量"章节下的正文段落
"""
import os, re, copy
from docx import Document
from docx.shared import Pt, Emu
from docx.oxml.ns import qn
from lxml import etree

_HERE     = os.path.dirname(os.path.abspath(__file__))
_ROOT     = os.path.dirname(os.path.dirname(_HERE))
TEMPLATE  = os.path.join(_ROOT, "assets", "营销专业基本情况 （2026年7月20日）(1).docx")


def _clear_runs(para):
    """清空段落所有 run，保留段落格式"""
    for run in para.runs:
        run.text = ""
    # 删除多余的 run XML 元素，只保留第一个
    p = para._p
    runs = p.findall(qn('w:r'))
    for r in runs[1:]:
        p.remove(r)


def _set_para_text(para, text):
    """清空段落，用单个 run 写入新文字，继承原有字体格式"""
    if not para.runs:
        return
    # 取原第一个 run 的格式作为模板
    ref_run = para.runs[0]
    _clear_runs(para)
    ref_run.text = text


def generate_marketing_docx(new_paragraphs: list[str],
                             base_docx_path: str = None,
                             output_path: str = None) -> str:
    """
    new_paragraphs: 新的电力电量正文段落列表（1段或2段）
    base_docx_path: 若有附件 docx，以此为基础；否则用模版
    output_path:    输出路径
    """
    src = base_docx_path if base_docx_path else TEMPLATE
    doc = Document(src)

    # 找到"一、电力电量"段落的位置
    elec_idx = None
    for i, para in enumerate(doc.paragraphs):
        if para.text.strip().startswith("一、电力电量") or para.text.strip() == "一、电力电量":
            elec_idx = i
            break

    if elec_idx is None:
        raise ValueError("未找到电力电量段落")

    # 找到该章节下的正文段落（到下一个章节标题为止）
    body_paras = []
    i = elec_idx + 1
    while i < len(doc.paragraphs):
        t = doc.paragraphs[i].text.strip()
        # 遇到下一章节标题停止
        if re.match(r'^[一二三四五六七八九十]+[、．.]', t):
            break
        if t:  # 非空段落
            body_paras.append(i)
        i += 1

    # 替换/增删正文段落
    if len(body_paras) == 0:
        raise ValueError("未找到电力电量正文段落")

    # 用新内容逐个替换现有段落
    for j, new_text in enumerate(new_paragraphs):
        if j < len(body_paras):
            _set_para_text(doc.paragraphs[body_paras[j]], new_text)
        else:
            # 需要在章节末尾新增段落，复制上一个段落的格式
            ref_para = doc.paragraphs[body_paras[-1]]
            new_para_xml = copy.deepcopy(ref_para._p)
            ref_para._p.addnext(new_para_xml)
            # 找到新插入的段落对象并设置文字
            # 重新找一下（段落列表已变）
            new_idx = list(doc.element.body).index(new_para_xml)
            from docx.text.paragraph import Paragraph
            np = Paragraph(new_para_xml, doc)
            _set_para_text(np, new_text)

    # 如果新内容段落数少于原有段落数，清空多余段落
    for j in range(len(new_paragraphs), len(body_paras)):
        _set_para_text(doc.paragraphs[body_paras[j]], "")

    if not output_path:
        output_path = os.path.join(_ROOT, "output", "marketing_output.docx")
    os.makedirs(os.path.dirname(output_path), exist_ok=True)
    doc.save(output_path)
    return output_path
